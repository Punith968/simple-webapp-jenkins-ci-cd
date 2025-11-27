from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os

PROJECT = "simple-webapp-jenkins-ci-cd"
AUTHOR = "Punith C"
DATE = datetime.now().strftime("%B %d, %Y")

# Known project facts
BUILD_NUM = "#10"
SUCCESS_RATE = "100% (4/4 stages)"
URL = "http://localhost:8090"
IMAGE_ID = "e54b76dbdf5f"
CONTAINER_NAME = "simple-webapp-demo"
CONTAINER_ID = "1db05a33b511"

SCREENSHOTS = [
    ("screenshots/jenkins_build.png", "Jenkins Build #10 — All stages green"),
    ("screenshots/docker_ps.png", "Docker container simple-webapp-demo (8090→80)"),
    ("screenshots/app_browser.png", "App at http://localhost:8090 with JS alert"),
]


def add_title(doc: Document, title: str, subtitle: str = ""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(22)
    if subtitle:
        p2 = doc.add_paragraph(subtitle)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.runs[0].font.size = Pt(12)


def add_h1(doc: Document, text: str):
    h = doc.add_heading(text, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_h2(doc: Document, text: str):
    h = doc.add_heading(text, level=2)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_bullets(doc: Document, items):
    for it in items:
        p = doc.add_paragraph(it, style='List Bullet')


def try_add_picture(doc: Document, path: str, caption: str):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(6))
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True
    else:
        p = doc.add_paragraph(f"[Add screenshot: {caption} — missing file {path}]")
        p.runs[0].italic = True


def build_document():
    doc = Document()

    # Title Page
    add_title(doc, "DevOps Case Study Report", f"Project: {PROJECT}  |  Author: {AUTHOR}  |  Date: {DATE}")
    doc.add_page_break()

    # Executive Summary / Abstract
    add_h1(doc, "Executive Summary")
    doc.add_paragraph(
        "This report presents a complete CI/CD pipeline for a static web application using Jenkins. "
        "The pipeline automates checkout, validation, testing, and Docker-based deployment. The final build "
        f"({BUILD_NUM}) achieved {SUCCESS_RATE} success and deployed a containerized app accessible at {URL}."
    )

    # Problem Statement & Objectives
    add_h1(doc, "Problem Statement & Objectives")
    doc.add_paragraph(
        "Create a CI/CD pipeline for a static HTML/CSS/JS website using Jenkins and implement automated "
        "deployment. Objectives included ensuring reproducible builds, minimal manual steps, and clear pipeline visibility."
    )
    add_h2(doc, "Objectives")
    add_bullets(doc, [
        "Automate Checkout → Build → Test → Deploy",
        "Validate required files before deployment",
        "Use Docker/Nginx for reproducible deployment",
        "Document results and lessons learned",
    ])

    # Scope
    add_h1(doc, "Scope")
    add_bullets(doc, [
        "Static website: index.html, styles.css, script.js",
        "Jenkins declarative pipeline (4 stages + post)",
        "Dockerfile for Nginx-based container",
        "Optional docker-compose for local Jenkins + webapp",
    ])

    # Environment & Architecture
    add_h1(doc, "Environment & Architecture")
    doc.add_paragraph(
        "Windows host with WSL2 Ubuntu as the Jenkins agent. Docker used for building and running the containerized web app. "
        "Architecture flow: Developer (GitHub) → Jenkins Pipeline → Docker Image → Nginx Container → Browser."
    )

    # Technologies and Tools
    add_h1(doc, "Technologies Used")
    add_bullets(doc, [
        "HTML5, CSS3, JavaScript",
        "Nginx (nginx:alpine)",
        "Docker",
        "Jenkins (Declarative Pipeline)",
        "Git/GitHub",
    ])

    add_h1(doc, "Tools Used")
    add_bullets(doc, [
        "Jenkins 2.528.2",
        "Docker 28.2.2",
        "Git 2.43.0",
        "Windows + WSL2 (Ubuntu 24.04)",
        "VS Code",
    ])

    # Methodology / Implementation Steps
    add_h1(doc, "Methodology / Implementation Steps")
    add_bullets(doc, [
        "Create app files and Dockerfile; initialize Git repo",
        "Configure Jenkins (Pipeline from SCM; set script path)",
        "Make scripts executable; add Jenkins to docker group; restart Jenkins",
        "Run validate.sh in Build & Test stages",
        "Build Docker image and run container on port 8090",
        "Post-deploy verification and logging",
    ])

    # Jenkins Pipeline
    add_h1(doc, "Jenkins Pipeline")
    add_bullets(doc, [
        "Stages: Checkout, Build (validate), Test (validate), Deploy (docker)",
        "Environment: DEPLOY_METHOD=docker; CONTAINER_NAME=simple-webapp-demo; WEBAPP_PORT=8090",
        "Nested repo path handled with dir() blocks",
    ])

    # Deployment
    add_h1(doc, "Deployment")
    add_bullets(doc, [
        "Image: simple-webapp:latest (Image ID: %s)" % IMAGE_ID,
        "Container: %s (ID: %s)" % (CONTAINER_NAME, CONTAINER_ID),
        "Port mapping: 8090 → 80; URL: %s" % URL,
        "Automated stop/remove old container before redeploy",
    ])

    # Results and Metrics
    add_h1(doc, "Results & Metrics")
    add_bullets(doc, [
        "Build %s: SUCCESS (4/4 stages)" % BUILD_NUM,
        "Total runtime ~15 seconds",
        "Application functional; JS alert verified",
    ])

    # Screenshots
    add_h1(doc, "Screenshots")
    for path, caption in SCREENSHOTS:
        try_add_picture(doc, path, caption)

    # Challenges & Resolutions
    add_h1(doc, "Challenges & Resolutions")
    add_bullets(doc, [
        "Jenkinsfile path + nested repo → set Script Path + dir()",
        "Script permissions on Linux → git chmod + pipeline chmod",
        "Docker permission denied → add Jenkins to docker group + restart",
        "Image pull timeout → pre-cache nginx:alpine",
        "Deploy method → switched from sudo copy to Docker",
    ])

    # Lessons Learned
    add_h1(doc, "Lessons Learned")
    add_bullets(doc, [
        "Docker simplifies CI/CD deployments",
        "Repo structure impacts Jenkins configuration",
        "Windows→Linux scripts need executable flags",
        "Group membership changes require service restart",
    ])

    # Conclusion
    add_h1(doc, "Conclusion")
    doc.add_paragraph(
        f"The CI/CD pipeline met all objectives, delivering {SUCCESS_RATE} success with a reproducible Docker-based deployment. "
        "The project is production-ready and demonstrates a clear, maintainable Jenkins workflow for static sites."
    )

    # References
    add_h1(doc, "References")
    add_bullets(doc, [
        "Repo: https://github.com/Punith968/simple-webapp-jenkins-ci-cd",
        "Jenkins Docs: https://www.jenkins.io/doc/",
        "Docker Docs: https://docs.docker.com/",
    ])

    # Appendix (Key Commands)
    add_h1(doc, "Appendix — Key Commands")
    doc.add_paragraph("docker build -t simple-webapp:latest .")
    doc.add_paragraph("docker run -d --name simple-webapp-demo -p 8090:80 simple-webapp:latest")
    doc.add_paragraph("./validate.sh")

    out_path = os.path.join(os.getcwd(), "project_report.docx")
    doc.save(out_path)
    print("Saved:", out_path)


if __name__ == "__main__":
    build_document()
